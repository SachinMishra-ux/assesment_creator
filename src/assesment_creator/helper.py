from ensure import ensure_annotations
import docx
import os

@ensure_annotations
def filter_text(text:str) -> tuple[str, list] :
    text_list = []

    text_list = text.split("\n")
    text_list = [t for t in text_list if "1 point" not in t and "5 points" not in t and t!="*"]

    return text_list[0], text_list[1:]

@ensure_annotations
def filter_qno(raw_qno:list) -> tuple[list, list]:
    questions = []
    options = []
    for text in raw_qno:
        q, a = filter_text(text)
        questions.append(q)
        options.append(a)

    return questions, options

@ensure_annotations
def create_docx(questions:list, options:list, output_file:str = "assessment.docx") -> None:

    doc = docx.Document()

    for question, options in zip(questions,options):
        doc.add_paragraph(question)
        for option in options:
            doc.add_paragraph(option)
        doc.add_paragraph("\n")

    os.makedirs("data", exist_ok=True)
    doc.save(os.path.join(os.path.curdir, "data", output_file))
