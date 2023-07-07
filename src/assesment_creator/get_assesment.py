from assesment_creator.logger import logger
from assesment_creator.cutom_exception import InvalidURLException
from assesment_creator.helper import output_file
from ensure import ensure_annotations
import re
import requests
from bs4 import BeautifulSoup
from time import sleep
import docx
import os
#url= "https://docs.google.com/forms/d/1WLfwaz1fuiueMGE8iEigwjZSP2HEUPGlyIWE9QNYhpg/edit"


@ensure_annotations
def get_data(link: str) -> list:
    page= requests.get(link)
    sleep(3)
    soup = BeautifulSoup(page.content, 'html.parser')
    content = soup.find_all(class_='geS5n')
    content= [pair for pair in content]
    content= str(content)
    mixed_pattern= 'M7eMe">(.*?)\<|auto">(.*?)\<\/span>'
    qno_list = re.findall(mixed_pattern, content)
    return qno_list

@ensure_annotations
def create_docx(qno_list:list, output_file:str = "assessment.docx") -> None:

    doc = docx.Document()
    #Add a Title to the document
    doc.add_heading('Assesment', 0)
    for i, qno in enumerate(qno_list):
        if len(qno[0])!=0:
            doc.add_paragraph(str(qno[0]))
        else:
            doc.add_paragraph(str(qno[1]))
    os.makedirs("data", exist_ok=True)
    doc.save(os.path.join(os.path.curdir, "data", output_file))


def create_assesment(form_link:str, file_name:str) -> None:
    try:
        if form_link and file_name is None:
            raise InvalidURLException("Url or file_name cannot be empty")
        else:
            output_file(file_name)
            content_list= get_data(form_link)
            if content_list:
                create_docx(content_list, file_name)
                print(f"Scraping done successfully! Check {file_name}\nThank you! \\m/")
            else:
                print("No data found! Please check the link provided")   
                
    except Exception as e:
        print(e)

